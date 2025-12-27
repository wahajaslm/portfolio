from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Create the Document
doc = Document()

# Set Margins (0.5 inches)
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Helper for Bottom Border
def add_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

# --- STYLE CONFIGURATION ---
# Base Text (DM Sans)
style_normal = doc.styles['Normal']
style_normal.font.name = 'DM Sans'
style_normal.font.size = Pt(10)
style_normal.paragraph_format.space_after = Pt(6)

# Helper to apply "Monument Extended" (Headings)
def apply_header_font(run, size=11, bold=True):
    run.font.name = "Monument Extended"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)

# Helper to apply "DM Sans" (Body)
def apply_body_font(run, size=10, bold=False, italic=False):
    run.font.name = "DM Sans"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

# --- HEADER SECTION ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("WAHAJ ASLAM")
apply_header_font(r, size=22, bold=True)
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Senior Audio DSP Engineer | Audio Codec Specialist | MPEG Audio Standards")
apply_body_font(r, size=11, bold=False)
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Nürnberg, Germany | wahajaslam08@gmail.com | +49 176 56968258 | linkedin.com/in/wahaj-aslam")
apply_body_font(r, size=10, bold=False)
p.paragraph_format.space_after = Pt(24) 

# --- SECTION HEADER GENERATOR ---
def add_section_head(text):
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    # Reverted to apply_header_font (Monument Extended)
    apply_header_font(r, size=12, bold=True)
    # Add Divider
    add_bottom_border(p)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True

# --- 1. PROFESSIONAL SUMMARY ---
add_section_head("PROFESSIONAL SUMMARY")
p = doc.add_paragraph("Senior Audio DSP Engineer with over eight years of experience developing advanced audio codecs for real-time and streaming applications. Strong background in audio signal processing and perceptual coding, with hands-on contributions to MPEG standards including xHE-AAC, MPEG-H 3D Audio, and AAC. Focused on turning research concepts into robust, production-quality C++ implementations, with emphasis on performance, audio quality, and long-term maintainability.")
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# --- 2. CORE TECHNICAL EXPERTISE ---
add_section_head("CORE TECHNICAL EXPERTISE")
skills = [
    ("Audio Codec & DSP", "Audio coding algorithms, perceptual models, speech & audio compression, bandwidth extension, source separation, STFT, LPC analysis, psychoacoustics, immersive audio"),
    ("Programming & Systems", "C, C++, Python, Bash, embedded Linux, multithreading, performance optimization, real-time systems"),
    ("Embedded Systems", "Embedded C/C++, ARM Cortex-M, DSP on embedded targets, memory-constrained systems, cross-compilation, hardware–software integration"),
    ("Frameworks & Tooling", "FFmpeg, Adobe Audition, sox, JUCE, VST, Windows Media Foundation, MATLAB, Pure Data, PortAudio, iOS CoreAudio, CMake, Git, GitLab CI/CD, Visual Studio, Xcode"),
    ("Python & Scientific Computing", "NumPy, SciPy, librosa, matplotlib, Jupyter"),
    ("AI / ML (Independent)", "PyTorch, Microsoft Agent Framework")
]

for cat, desc in skills:
    p = doc.add_paragraph(style='List Bullet')
    r_cat = p.add_run(cat + ": ")
    apply_body_font(r_cat, bold=True)
    r_desc = p.add_run(desc)
    apply_body_font(r_desc)
    p.paragraph_format.space_after = Pt(2) 

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# --- 3. PROFESSIONAL EXPERIENCE ---
add_section_head("PROFESSIONAL EXPERIENCE")

def add_job_entry(title, company, dates, bullets):
    # Job Title 
    # Job Title + Date (Same Line)
    p_title = doc.add_paragraph()
    # Add Right Tab Stop for Date
    p_title.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
    
    r_title = p_title.add_run(title)
    # Job Title -> DM Sans
    apply_body_font(r_title, size=10, bold=True)
    
    # Tab -> Date
    r_tab = p_title.add_run("\t")
    r_date = p_title.add_run(dates)
    apply_body_font(r_date, italic=True)

    p_title.paragraph_format.space_after = Pt(2)
    p_title.paragraph_format.keep_with_next = True

    # Company Name (New Line)
    p_comp = doc.add_paragraph()
    r_comp = p_comp.add_run(company)
    apply_body_font(r_comp, italic=True)
    p_comp.paragraph_format.space_after = Pt(6) # Adjusted space
    p_comp.paragraph_format.keep_with_next = True

    # Bullets
    for b in bullets:
        p_b = doc.add_paragraph(style='List Bullet')
        r_b = p_b.add_run(b)
        apply_body_font(r_b)
        p_b.paragraph_format.space_after = Pt(2)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

# Job 1
add_job_entry(
    "Senior Engineer – Audio Codec Development",
    "Fraunhofer Institute for Integrated Circuits IIS, Erlangen",
    "Apr 2018 – Present",
    [
        "Development and maintenance of core audio coding components in xHE-AAC, MPEG-H 3D Audio, and AAC encoders",
        "Lead roles for design and implementation of new codec tools from early research through production deployment",
        "Architectural contributions related to low-latency processing, adaptive streaming, and scalability",
        "Performance and complexity optimization on ARM and x86_64 platforms",
        "Integration and validation within FFmpeg and Windows Media Foundation pipelines",
        "Long-term maintenance, testing, and regression analysis for production codebases"
    ]
)

# Job 2
add_job_entry(
    "Scientific Researcher – Audio Signal Processing",
    "Fraunhofer IIS, Erlangen",
    "Aug 2016 – Sep 2017",
    [
        "Research on audio and speech coding algorithms including bandwidth extension and signal reconstruction",
        "Algorithm design and evaluation using MATLAB/Python and C",
        "Objective and subjective evaluation using MUSHRA listening tests",
        "Collaboration with engineering teams to assess feasibility for production integration"
    ]
)

# Job 3
add_job_entry(
    "Audio DSP Engineer – Research Internship",
    "Fraunhofer IIS, Erlangen",
    "Dec 2016 – Mar 2017",
    [
        "Implementation of MPEG-H audio coding tools in C and MATLAB",
        "Development of multichannel coding tools using time-differential techniques",
        "Evaluation of bitrate efficiency and perceptual quality"
    ]
)

# Job 4
add_job_entry(
    "Software Engineer – Embedded LTE/4G Protocol Stack",
    "u-blox, Pakistan",
    "Jul 2012 – Sep 2014",
    [
        "Development of LTE/4G NAS-layer components according to 3GPP Releases 9–11",
        "Implementation of AT command handling and USIM modules in embedded C",
        "Customization of protocol stack components for embedded devices",
        "Protocol verification using Anite conformance tools"
    ]
)

# --- 4. EDUCATION ---
add_section_head("EDUCATION")

p = doc.add_paragraph()
r1 = p.add_run("M.Sc. Information & Communication Engineering")
apply_body_font(r1, bold=True)
p.add_run("\n")
r2 = p.add_run("Technical University of Darmstadt, Germany")
apply_body_font(r2, italic=True)
p.paragraph_format.space_after = Pt(8)

p = doc.add_paragraph()
r3 = p.add_run("B.Sc. Telecommunication Engineering")
apply_body_font(r3, bold=True)
p.add_run("\n")
r4 = p.add_run("National University of Computer & Emerging Sciences (FAST-NU), Pakistan")
apply_body_font(r4, italic=True)
p.paragraph_format.space_after = Pt(12)

# --- 5. PATENTS ---
add_section_head("PATENTS")
p = doc.add_paragraph(style='List Bullet')
r_pat = p.add_run("WO2023021137A1 (Granted & Published)")
apply_body_font(r_pat, bold=True)
p.add_run(" – Audio encoder and method for providing an encoded representation of audio information")

# --- 6. LANGUAGES ---
add_section_head("LANGUAGES")
doc.add_paragraph("English – Fluent | German – B1 | Urdu – Native")

# --- 7. INTERESTS ---
add_section_head("INTERESTS")
doc.add_paragraph("Audio signal processing beyond production work, exploratory use of machine learning for audio analysis, Team sports, cycling, hiking, and interest in clean engineering focused tooling.")

# Save
doc.save("Wahaj_Aslam_CV_Final3.docx")
print("Done! File saved as Wahaj_Aslam_CV_Final3docx")